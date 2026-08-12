from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "价策AI产品使用说明书.docx"
ASSETS = ROOT / "docs" / "assets"

NAVY = RGBColor(31, 41, 64)
BLUE = RGBColor(79, 111, 232)
MUTED = RGBColor(110, 120, 138)
LIGHT = "F3F6FD"
GRID = "D9DEE9"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_end)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        body = paragraph.add_run(text[len(bold_lead):])
        set_run_font(body)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        set_run_font(paragraph.add_run(item))


def add_steps(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        set_run_font(paragraph.add_run(item))


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(cell.paragraphs[0].add_run(header), size=9.5, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            paragraph = cells[idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph.add_run(str(value)), size=9.2)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc, filename, caption):
    path = ASSETS / filename
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    picture = run.add_picture(str(path), width=Inches(6.25))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", filename)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(8)
    set_run_font(cap.add_run(caption), size=8.5, color=MUTED, italic=True)


def new_content_page(doc):
    doc.add_page_break()


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run("价策 AI  |  产品使用说明书"), size=8.5, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])

    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_run_font(kicker.add_run("PRICE SCOPE  ·  OPERATOR GUIDE"), size=10, color=BLUE, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(9)
    set_run_font(title.add_run("价策 AI"), size=31, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    set_run_font(subtitle.add_run("智能定价决策平台 · 产品使用说明书"), size=15, color=BLUE, bold=True)
    description = doc.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.space_after = Pt(70)
    set_run_font(description.add_run("面向电商运营、品类负责人、店铺负责人、采购与定价人员"), size=10, color=MUTED)
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(version.add_run("版本 1.2  |  价格拐点分析版"), size=11, color=NAVY, bold=True)
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(date.add_run("2026 年 8 月"), size=9.5, color=MUTED)

    new_content_page(doc)
    add_heading(doc, "目录与快速导航", 1)
    for index, item in enumerate((
        "产品简介与快速开始", "主界面与经营驾驶舱", "价格采集中心", "数据导入",
        "价格诊断", "价格拐点分析", "定价策略体系", "调价模拟器", "报告中心",
        "建议操作流程", "数据安全与常见问题", "后续商用升级",
    ), 1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(7)
        set_run_font(paragraph.add_run(f"{index:02d}"), size=10, color=BLUE, bold=True)
        set_run_font(paragraph.add_run(f"    {item}"), size=11, color=NAVY)
    add_heading(doc, "1. 产品简介", 1)
    add_body(doc, "价策 AI 是面向电商商家的价格情报与定价决策产品。系统从平台开放 API、合规公开商品页或 CSV 获得竞品报价，结合商品价格、成本、销量、库存与商品角色，完成价格诊断、建议区间、调价模拟和可视化报告。")
    add_body(doc, "当前版本不会自动修改平台售价。演示业务数据保存在浏览器；本地采集服务只监听本机地址，不绕过登录、验证码、robots.txt 或访问控制。")
    add_heading(doc, "2. 快速开始", 1)
    add_steps(doc, [
        "安装 Node.js 22.13 或更高版本。",
        "进入 price-scope-ai 目录，依次执行 pnpm install 和 pnpm dev。",
        "打开终端显示的本地地址，默认通常为 http://127.0.0.1:4173。",
        "首次进入时，系统自动加载 12 个母婴奶粉演示商品。",
        "如需采集公开页，另开终端执行 pnpm collector:serve。",
    ])

    new_content_page(doc)
    add_heading(doc, "3. 主界面与经营驾驶舱", 1)
    add_figure(doc, "dashboard.png", "图 1  经营驾驶舱：全店价格健康度和优先调价机会")
    add_body(doc, "左侧导航包含经营驾驶舱、价格采集中心、价格诊断、价格拐点分析、调价模拟器、报告中心、数据中心和使用指引八个模块。")
    add_bullets(doc, [
        "近 30 天销售额、预计增量机会、综合毛利额与竞品覆盖率。",
        "价格竞争力矩阵、价格健康度分布与优先调价商品。",
        "核心商品的本店价格与市场价格趋势。",
    ])
    add_body(doc, "价格竞争力矩阵中，横轴表示价格指数，纵轴表示近 30 天销量，气泡大小代表毛利空间。商品越靠右，说明本店价格相对市场越高。")

    new_content_page(doc)
    add_heading(doc, "4. 价格采集中心", 1)
    add_body(doc, "采集中心将面试演示与真实采集明确隔离。演示快照无需平台密钥且标识为 DEMO；本地采集服务接收用户提供的公开商品页 HTTPS 链接，输出平台、价格、解析方式、采集时间和失败原因。")
    add_heading(doc, "4.1 本地采集步骤", 2)
    add_steps(doc, [
        "在项目目录运行 pnpm collector:serve。",
        "进入价格采集中心，切换为“本地采集服务”。",
        "粘贴公开商品页链接，每行一个，然后点击“立即采集”。",
        "核对采集状态、解析方式和匹配商品；失败项改用开放 API 或 CSV。",
    ])
    add_heading(doc, "4.2 支持范围与边界", 2)
    add_bullets(doc, [
        "支持京东、天猫、淘宝、拼多多、苏宁易购和唯品会 HTTPS 域名白名单。",
        "最小请求间隔 1.2 秒，检查 robots.txt，限制单次任务和页面大小。",
        "依次解析 JSON-LD、价格 Meta 和页面公开数据。",
        "不绕过登录、验证码或访问控制；商业部署优先使用平台开放 API。",
        "密钥、Cookie 和账号信息不得提交到 GitHub 或下发浏览器。",
    ])

    new_content_page(doc)
    add_heading(doc, "5. 数据导入", 1)
    add_body(doc, "进入“数据中心”，点击“下载导入模板”获取标准 CSV 模板。填写后点击“选择 CSV 文件”完成导入。")
    add_heading(doc, "5.1 必填字段", 2)
    add_table(doc, ["字段", "用途", "示例"], [
        ("商品编码", "SKU 唯一标识", "P001"), ("品牌", "商品标准化", "伊利"),
        ("系列", "同系列匹配", "金领冠珍护"), ("段位", "母婴商品分组", "2段"),
        ("规格", "可比规格判断", "750g"), ("当前价格", "计算价格指数", "289"),
        ("成本", "计算利润安全线", "209"), ("近30天销量", "策略模拟", "156"),
        ("商品角色", "差异化定价", "转化型"), ("平台价格", "市场价格样本", "261/257/251"),
    ], [1900, 4600, 2860])
    add_heading(doc, "5.2 导入要求", 2)
    add_bullets(doc, [
        "文件格式必须为 CSV，推荐 UTF-8 编码。",
        "价格字段只填写数字，不包含“元”或货币符号。",
        "转化率填写百分比数值，例如 4.2 表示 4.2%。",
        "平台价格应为用户实际可获得的到手价。",
        "每个商品建议至少提供 3 个竞品报价。",
        "导入新文件会替换当前浏览器中的商品数据。",
    ])

    new_content_page(doc)
    add_heading(doc, "6. 价格诊断", 1)
    add_figure(doc, "diagnosis.png", "图 2  商品级价格诊断与建议价格区间")
    add_body(doc, "价格诊断表包含当前到手价、市场中位价、价格指数、毛利率、诊断状态、建议价格区间和匹配可信度。点击“详情”可查看竞品样本和 AI 解释。")
    add_heading(doc, "6.1 诊断状态", 2)
    add_table(doc, ["状态", "判定", "处理建议"], [
        ("偏高价格", "价格指数 > 106", "优先进入调价模拟"),
        ("轻度偏高", "102 < 价格指数 ≤ 106", "结合经营角色审核"),
        ("有竞争力", "接近市场主流区间", "保持并持续监控"),
        ("优势价格", "低于主流且利润安全", "加强活动曝光"),
        ("低价风险", "明显偏低或毛利不安全", "检查成本和活动"),
    ], [1900, 3100, 4360])
    add_heading(doc, "6.2 关键公式", 2)
    add_bullets(doc, [
        "到手价 = 页面价格 + 运费 - 优惠券 - 满减 - 平台补贴。",
        "价格指数 = 本店到手价 ÷ 市场中位价 × 100。",
        "利润底价 = 商品成本 ÷（1 - 最低毛利率）。",
        "建议价格同时考虑市场中位价、商品角色与利润底价。",
    ])

    new_content_page(doc)
    add_heading(doc, "7. 价格拐点分析", 1)
    add_body(doc, "价格拐点分析比较连续价格的变化方向，只在斜率反转且价格变动超过敏感度阈值时生成信号，并联合销量响应、市场价差和可信度给出建议动作。")
    add_table(doc, ["信号", "含义", "建议动作"], [
        ("降价止跌拐点", "价格先下降后回升", "检查销量响应，维持并观察 3—7 天"),
        ("涨价转弱拐点", "价格先上涨后回落", "核对市场价差后进入调价模拟"),
        ("无显著拐点", "波动未超过阈值", "保持价格并继续采集历史数据"),
    ], [2200, 3100, 4060])
    add_bullets(doc, [
        "高敏感（1%）适合日常监控，标准（1.8%）适合周度运营，稳健（3%）适合管理层复盘。",
        "拐点是相关性信号，不代表调价必然造成销量变化。",
        "执行前应排除大促、缺货、广告投放和竞品规格变化。",
        "建议从拐点页面进入调价模拟，设置利润安全线后再做小范围实验。",
    ])

    new_content_page(doc)
    add_heading(doc, "8. 定价策略体系", 1)
    add_body(doc, "定价不是简单跟随最低价。系统先统一到手价口径，再结合市场价格带、利润底价和商品经营角色，生成进攻、平衡与保守方案。")
    add_heading(doc, "8.1 商品角色策略", 2)
    add_table(doc, ["角色", "经营目标", "建议动作"], [
        ("引流型", "获取访问与新客", "靠近市场低位，不击穿利润底价"),
        ("转化型", "提升下单效率", "中位价附近，先做 7 天实验"),
        ("利润型", "贡献毛利", "允许合理溢价，减少直接折扣"),
        ("防御型", "对标重点竞品", "紧跟目标竞品有效到手价"),
        ("形象型", "强化品牌定位", "跟随标杆，减少频繁波动"),
        ("清仓型", "加快库存周转", "利润与渠道约束内积极降价"),
    ], [1600, 3000, 4760])
    add_heading(doc, "8.2 策略护栏", 2)
    add_bullets(doc, [
        "匹配可信度不足时先人工审核，不给自动执行动作。",
        "建议价格不得低于利润底价。",
        "进攻价取建议区间低位，平衡价取中点，保守价取高位。",
        "促销价与日常价分开观察，组合装先换算单位价格。",
        "单次大幅调价拆成小流量实验，设置销量和毛利双重停止条件。",
    ])

    new_content_page(doc)
    add_heading(doc, "9. 调价模拟器", 1)
    add_figure(doc, "simulator.png", "图 3  调价前的销量、销售额与毛利模拟")
    add_steps(doc, [
        "选择需要模拟的商品。",
        "输入模拟价格，或选择进攻、平衡、保守方案。",
        "查看预计销量、销售额、毛利额和毛利率。",
        "检查系统是否提示“满足利润安全线”。",
        "将通过安全线的方案用于 7 天小范围实验。",
    ])
    add_body(doc, "当前版本使用固定价格弹性系数进行演示，模拟结果只用于辅助决策，不应代替实际业务实验。")

    new_content_page(doc)
    add_heading(doc, "10. 报告中心", 1)
    add_figure(doc, "reports.png", "图 4  面向管理层的价格竞争力诊断周报")
    add_bullets(doc, [
        "导出调价清单：下载 CSV，供运营人员批量审核和执行。",
        "打印 / 保存 PDF：调用浏览器打印功能，生成管理层周报。",
    ])
    add_body(doc, "管理层报告包含综合健康度、诊断商品数、重点优化商品、增量机会、数据覆盖率、经营结论和优先调价表。")

    new_content_page(doc)
    add_heading(doc, "11. 建议操作流程", 1)
    add_steps(doc, [
        "定期运行授权 API、公开页采集或 CSV 导入。",
        "检查采集成功率、采集时间和同款匹配可信度。",
        "查看最新价格拐点，核对市场价差与销量响应。",
        "优先处理偏高价格和低价风险商品。",
        "检查商品匹配和竞品样本是否合理。",
        "在模拟器中比较进攻、平衡与保守方案。",
        "导出调价清单，完成运营和负责人审核。",
        "对少量商品执行 7 天调价实验。",
        "复盘销量、转化率、销售额和毛利变化。",
        "验证有效后再扩大执行范围。",
    ])
    new_content_page(doc)
    add_heading(doc, "12. 数据安全与版本边界", 1)
    add_bullets(doc, [
        "当前数据保存在浏览器 localStorage 中，清除浏览器数据可能导致导入内容丢失。",
        "本地采集服务只监听 127.0.0.1，不对公网开放。",
        "系统不保存平台密钥、Cookie 或个人账号，也不会自动执行调价。",
        "不采集登录后、验证码保护或 robots.txt 禁止的页面。",
        "演示竞品价格不代表真实市场报价。",
        "真实商用前，应接入合规的数据接口、账号权限、审批记录和服务端数据库。",
        "AI 文字说明来自结构化计算结果，最终定价仍需要业务人员审核。",
    ])
    add_heading(doc, "13. 常见问题", 1)
    for question, answer in (
        ("公开页采集失败", "平台可能采用动态渲染、登录校验或 robots 限制；系统不会绕过控制，请改用官方开放 API 或 CSV。"),
        ("导入后没有商品", "检查 CSV 是否包含“当前价格”字段，并确认价格为大于 0 的数字。"),
        ("商品全部显示偏高", "检查竞品价格是否为到手价、规格是否一致，以及是否混合单件价和组合装总价。"),
        ("建议价格高于市场价格", "通常是利润底价或商品角色限制导致，应检查成本、最低毛利率和商品角色。"),
        ("如何生成 PDF", "进入报告中心，点击“打印 / 保存 PDF”，在浏览器打印窗口选择“另存为 PDF”。"),
        ("如何恢复演示数据", "进入数据中心，点击“恢复演示数据”。"),
    ):
        add_body(doc, f"{question}：{answer}", bold_lead=f"{question}：")
    add_heading(doc, "14. 后续商用升级建议", 1)
    add_bullets(doc, [
        "接入真实平台授权 API、授权续期和调用监控。",
        "增加定时任务、历史报价数据库与异常告警。",
        "增加多商家账号、权限与数据隔离。",
        "增加商品匹配人工审核工作台。",
        "保存调价审批、执行与实验记录。",
        "使用真实历史数据训练品类价格弹性模型。",
        "增加邮件、企业微信或钉钉定时报告。",
        "在完整权限与安全机制下接入平台调价执行。",
    ])

    core = doc.core_properties
    core.title = "价策 AI 产品使用说明书"
    core.subject = "价格采集与智能定价决策平台操作指南"
    core.author = "价策 AI 产品团队"
    core.keywords = "电商, 定价, 比价, AI, 产品说明书"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
